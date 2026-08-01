#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""docx 段落文本替换，保留 run 格式。

用法:
    python3 docx_replace_text.py <file.docx> <old_text> <new_text> [--all]
    --all: 替换所有匹配段落；默认只替换第一个匹配段落

原则:
    - 优先在 run 内部做替换（保留该 run 的字体/字号/颜色）
    - 若 old_text 跨多个 run，先合并段落文本做替换，再用第一个 run 的格式重建
    - 重建前检查 runs 数量；多 run 且格式不同时谨慎处理

来源: 《World War C》翻译项目复盘教训 #7 (set_text 整段替换丢格式)
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def replace_in_paragraph(p, old, new):
    """在段落内替换文本，尽量保留 run 格式。返回是否替换成功。"""
    # 先尝试在单个 run 内替换（最安全，格式完全保留）
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            return True
    # 跨 run 情况：整段重建，用第一个 run 的格式
    full = ''.join(r.text for r in p.runs)
    if old in full:
        fmt = p.runs[0] if p.runs else None
        size = fmt.font.size if fmt else None
        name = fmt.font.name if fmt else None
        bold = fmt.font.bold if fmt else None
        color = fmt.font.color.rgb if fmt and fmt.font.color and fmt.font.color.type is not None else None
        east = None
        if fmt is not None and fmt._element.rPr is not None and fmt._element.rPr.rFonts is not None:
            east = fmt._element.rPr.rFonts.get(qn('w:eastAsia'))
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        r = p.add_run(full.replace(old, new))
        if size: r.font.size = size
        if name: r.font.name = name
        if bold is not None: r.font.bold = bold
        if color is not None: r.font.color.rgb = color
        rPr2 = r._element.get_or_add_rPr()
        rFonts = rPr2.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts'); rPr2.append(rFonts)
        if east: rFonts.set(qn('w:eastAsia'), east)
        return True
    return False


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    path, old, new = sys.argv[1], sys.argv[2], sys.argv[3]
    replace_all = '--all' in sys.argv
    doc = Document(path)
    count = 0
    for p in doc.paragraphs:
        if replace_in_paragraph(p, old, new):
            count += 1
            if not replace_all:
                break
    doc.save(path)
    print(f'✅ 替换 {count} 处: "{old}" → "{new}"')


if __name__ == '__main__':
    main()
